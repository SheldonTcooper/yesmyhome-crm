#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def add_horizontal_line(paragraph):
    """Add a horizontal line"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4A90E2')
    pBdr.append(bottom)
    pPr.append(pBdr)

doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ============ CAPA ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("ASSESSMENT TÉCNICO")
title_run.font.size = Pt(36)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(88, 166, 255)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("CRM YesMyHome + n8n + Remarketing")
subtitle_run.font.size = Pt(16)
subtitle.paragraph_format.space_before = Pt(6)

doc.add_paragraph()

# INFO TABLE
info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Table Grid'
info_table.rows[0].cells[0].text = "Data"
info_table.rows[0].cells[1].text = "30 de julho de 2026"
info_table.rows[1].cells[0].text = "Cliente"
info_table.rows[1].cells[1].text = "YesMyHome Negociações Imobiliárias"
info_table.rows[2].cells[0].text = "Escopo"
info_table.rows[2].cells[1].text = "Personalização + Integração + Automação"
info_table.rows[3].cells[0].text = "Timeline"
info_table.rows[3].cells[1].text = "3-4 semanas"
info_table.rows[4].cells[0].text = "Status"
info_table.rows[4].cells[1].text = "Pronto para Implementação"

doc.add_page_break()

# ============ ÍNDICE ============
doc.add_heading("ÍNDICE", level=1)
toc_items = [
    "1. Visão Geral",
    "2. Personalização do CRM Laravel",
    "3. Integração com Site Jetimob",
    "4. Arquitetura n8n para Remarketing",
    "5. Fluxos de Automação Detalhados",
    "6. Timeline de Implementação",
    "7. Riscos e Mitigação",
    "8. Próximos Passos",
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============ 1. VISÃO GERAL ============
doc.add_heading("1. VISÃO GERAL", level=1)

doc.add_heading("1.1 Objetivo", level=2)
doc.add_paragraph(
    "Implementar um sistema integrado de gestão de leads imobiliários que:"
)
objetivos = [
    "Captura automática de leads do site yesmyhome.com.br (Jetimob)",
    "Centraliza todos os leads em um CRM Laravel personalizável",
    "Envia mensagens automáticas (SMS, Email, WhatsApp)",
    "Realiza análise inteligente com Claude IA",
    "Executa remarketing automático via n8n",
    "Aumenta conversão com follow-ups programados",
]
for obj in objetivos:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading("1.2 Arquitetura Geral", level=2)
arch_para = doc.add_paragraph()
arch_para.add_run("Fluxo Completo:\n").bold = True
doc.add_paragraph("Site Jetimob → CRM Laravel → n8n → Claude AI → WhatsApp/Email/SMS", style='List Bullet')

doc.add_heading("1.3 Tecnologias", level=2)
techs = [
    "Backend: Laravel 12 (PHP 8.5.9)",
    "Frontend: Blade Templates + Bootstrap 5",
    "Banco: SQLite (desenvolvimento) / PostgreSQL (produção)",
    "Automação: n8n (self-hosted)",
    "IA: Claude API (Haiku/Sonnet)",
    "Mensageria: WhatsApp Business API Official",
    "Email: Resend",
    "Node.js: Backend de processamento (localhost:5000)",
]
for tech in techs:
    doc.add_paragraph(tech, style='List Bullet')

doc.add_page_break()

# ============ 2. PERSONALIZAÇÃO DO CRM ============
doc.add_heading("2. PERSONALIZAÇÃO DO CRM LARAVEL", level=1)

doc.add_heading("2.1 Estrutura Base (Krayin)", level=2)
doc.add_paragraph(
    "O CRM utiliza o Krayin, um framework Laravel baseado em pacotes que permite "
    "modular e extensível."
)

doc.add_heading("2.2 Customizações Necessárias", level=2)

# 2.2.1
doc.add_heading("2.2.1 Modelo de Dados - Leads", level=3)
doc.add_paragraph("Campos principais a serem adicionados/customizados:")
lead_fields = [
    "name (nome completo)",
    "email (email do cliente)",
    "phone (telefone com país)",
    "whatsapp (WhatsApp com país)",
    "property_type (tipo de imóvel: apartamento, casa, terreno, etc)",
    "operation_type (tipo de operação: compra, aluguel, venda)",
    "budget (orçamento em reais)",
    "location (localização desejada)",
    "notes (observações internas)",
    "source (origem do lead: site, telefone, whatsapp, etc)",
    "status (novo, conversando, visitando, proposta, fechado)",
    "ai_score (score de 0-100 da IA)",
    "ai_analysis (análise textual da IA)",
    "assigned_to (agente responsável)",
    "last_contact (última interação)",
    "next_followup (próximo follow-up agendado)",
]
for field in lead_fields:
    doc.add_paragraph(field, style='List Bullet')

# 2.2.2
doc.add_heading("2.2.2 Customizações de Interface", level=3)

customizacoes_ui = {
    "Dashboard": [
        "Gráfico de leads por status (novo/conversando/fechado)",
        "KPIs: total de leads, conversão %, segups pendentes",
        "Últimos leads criados (tempo real)",
        "Score médio dos leads",
    ],
    "Listagem de Leads": [
        "Filtros avançados (status, tipo, orçamento, data)",
        "Busca por nome/email/telefone",
        "Bulk actions (mudar status, atribuir, etc)",
        "Exportar para Excel/CSV",
    ],
    "Detalhe do Lead": [
        "Timeline completa de interações",
        "Análise IA com badges de propensão",
        "Histórico de mensagens (SMS/Email/WhatsApp)",
        "Notas internas com timestamp",
        "Agendador de follow-up",
    ],
}

for section, items in customizacoes_ui.items():
    p = doc.add_paragraph()
    p.add_run(section + ": ").bold = True
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# 2.2.3
doc.add_heading("2.2.3 Controllers e Rotas", level=3)

doc.add_paragraph("Rotas API a serem criadas/modificadas:")

rotas_table = doc.add_table(rows=9, cols=3)
rotas_table.style = 'Table Grid'
rotas_table.rows[0].cells[0].text = "Método"
rotas_table.rows[0].cells[1].text = "Rota"
rotas_table.rows[0].cells[2].text = "Função"

rotas_data = [
    ("POST", "/api/webhook/create-lead", "Recebe lead do Jetimob"),
    ("GET", "/api/leads", "Lista todos os leads"),
    ("GET", "/api/leads/{id}", "Detalhes do lead"),
    ("PUT", "/api/leads/{id}", "Atualiza lead"),
    ("POST", "/api/leads/{id}/message", "Envia mensagem manual"),
    ("POST", "/api/leads/{id}/followup", "Agenda follow-up"),
    ("GET", "/api/analytics", "Retorna métricas"),
    ("POST", "/api/webhook/n8n", "Recebe resposta do n8n"),
]

for idx, (metodo, rota, funcao) in enumerate(rotas_data, start=1):
    rotas_table.rows[idx].cells[0].text = metodo
    rotas_table.rows[idx].cells[1].text = rota
    rotas_table.rows[idx].cells[2].text = funcao

# 2.2.4
doc.add_heading("2.2.4 Modelos Eloquent", level=3)

models_info = {
    "Lead": "Modelo principal com relações para Messages, FollowUps, AIAnalysis",
    "Message": "Registra SMS, Email, WhatsApp enviados",
    "FollowUp": "Agendamentos de follow-ups",
    "AIAnalysis": "Armazena análise completa da Claude",
    "User": "Agentes/equipe com permissões",
    "LeadTag": "Tags para categorizar leads",
}

for model, desc in models_info.items():
    p = doc.add_paragraph()
    p.add_run(model + ": ").bold = True
    p.add_run(desc)

doc.add_page_break()

# ============ 3. INTEGRAÇÃO COM SITE JETIMOB ============
doc.add_heading("3. INTEGRAÇÃO COM SITE JETIMOB", level=1)

doc.add_heading("3.1 Arquitetura de Integração", level=2)

fluxo = [
    "Cliente preenche formulário no site (yesmyhome.com.br)",
    "Formulário faz POST para nosso endpoint webhook",
    "CRM recebe dados e cria lead no banco",
    "Dispara automação n8n via webhook",
    "N8n envia mensagens automáticas",
    "Lead aparece no dashboard do CRM",
]

for idx, passo in enumerate(fluxo, 1):
    p = doc.add_paragraph()
    p.add_run(f"Passo {idx}: ").bold = True
    p.add_run(passo)

doc.add_heading("3.2 Endpoint Webhook", level=2)

doc.add_paragraph("URL do Webhook (em produção):")
webhook_url = doc.add_paragraph()
webhook_run = webhook_url.add_run("https://seu-dominio.com/api/webhook/create-lead")
webhook_run.font.name = 'Courier New'
webhook_run.font.size = Pt(10)

doc.add_paragraph()

doc.add_paragraph("Método: POST")
doc.add_paragraph("Headers esperados:")
doc.add_paragraph("Content-Type: application/json", style='List Bullet')

doc.add_paragraph()
doc.add_paragraph("Body esperado (exemplo):")
payload = doc.add_paragraph()
payload_run = payload.add_run("""{
  "name": "João Silva",
  "email": "joao@email.com",
  "phone": "+5541999999999",
  "whatsapp": "+5541999999999",
  "property_type": "apartment",
  "operation_type": "sale",
  "budget": 500000,
  "location": "Curitiba, PR",
  "notes": "Interessado em apartamento 3 quartos"
}""")
payload_run.font.name = 'Courier New'
payload_run.font.size = Pt(9)

doc.add_heading("3.3 Como Configurar no Jetimob", level=2)

jetimob_steps = [
    "Acessar painel do Jetimob (admin area)",
    "Ir em Integrações ou Webhooks",
    "Criar novo webhook",
    "URL: https://seu-dominio.com/api/webhook/create-lead",
    "Evento: Quando formulário for submetido",
    "Método: POST",
    "Mapear campos do formulário Jetimob para nosso JSON",
    "Testar webhook (enviar lead de teste)",
    "Ativar webhook em produção",
]

for step in jetimob_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading("3.4 Mapeamento de Campos Jetimob → CRM", level=2)

mapping_table = doc.add_table(rows=9, cols=3)
mapping_table.style = 'Table Grid'
mapping_table.rows[0].cells[0].text = "Campo Jetimob"
mapping_table.rows[0].cells[1].text = "Campo CRM"
mapping_table.rows[0].cells[2].text = "Tipo"

mapping_data = [
    ("Nome", "name", "String"),
    ("E-mail", "email", "Email"),
    ("Telefone", "phone", "String"),
    ("WhatsApp", "whatsapp", "String"),
    ("Tipo de Imóvel", "property_type", "String"),
    ("Tipo de Operação", "operation_type", "String"),
    ("Mensagem", "notes", "Text"),
    ("Data/Hora", "created_at", "DateTime"),
]

for idx, (jetimob, crm, tipo) in enumerate(mapping_data, start=1):
    mapping_table.rows[idx].cells[0].text = jetimob
    mapping_table.rows[idx].cells[1].text = crm
    mapping_table.rows[idx].cells[2].text = tipo

doc.add_page_break()

# ============ 4. ARQUITETURA N8N PARA REMARKETING ============
doc.add_heading("4. ARQUITETURA N8N PARA REMARKETING", level=1)

doc.add_heading("4.1 O Que é n8n?", level=2)

doc.add_paragraph(
    "n8n é uma plataforma de automação open-source que permite criar workflows "
    "complexos sem código. Será instalado em seu servidor (self-hosted) e acionado "
    "por webhooks do CRM."
)

doc.add_heading("4.2 Instalação e Configuração", level=2)

instalacao = [
    "Opção: Docker Compose (recomendado)",
    "Porta: 5678 (acessível via n8n.seu-dominio.com)",
    "Autenticação: Habilitada",
    "Backups: Automáticos diários",
    "Storage: Volume persistente (/home/node/.n8n)",
]

for item in instalacao:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("4.3 Conceitos de n8n", level=2)

conceitos = {
    "Workflow": "Fluxo de automação completo (ex: welcome sequence)",
    "Trigger": "O que inicia o workflow (webhook do CRM, agendamento, etc)",
    "Node": "Ação individual (enviar email, fazer API call, etc)",
    "Connection": "Credenciais integradas (Claude API, Resend, etc)",
}

for conceito, desc in conceitos.items():
    p = doc.add_paragraph()
    p.add_run(conceito + ": ").bold = True
    p.add_run(desc)

doc.add_heading("4.4 Workflows Principais", level=2)

workflows = [
    "WF-01: Welcome Sequence (novo lead)",
    "WF-02: 24h Follow-up (sem resposta)",
    "WF-03: 48h Follow-up (sem resposta)",
    "WF-04: 72h Follow-up (sem resposta)",
    "WF-05: Lead Inativo 7 dias (remarketing)",
    "WF-06: Lead Qualificado (score > 80)",
    "WF-07: Relatório Diário (9h da manhã)",
]

for wf in workflows:
    doc.add_paragraph(wf, style='List Bullet')

doc.add_page_break()

# ============ 5. FLUXOS DE AUTOMAÇÃO DETALHADOS ============
doc.add_heading("5. FLUXOS DE AUTOMAÇÃO DETALHADOS", level=1)

# WF-01
doc.add_heading("5.1 WF-01: Welcome Sequence (Novo Lead)", level=2)

doc.add_paragraph("Acionado quando: Lead criado no CRM")
doc.add_paragraph()

ws_steps = [
    {
        "num": 1,
        "nome": "Receber Webhook",
        "desc": "CRM dispara webhook com dados do lead",
        "saida": "name, email, phone, whatsapp, property_type",
    },
    {
        "num": 2,
        "nome": "Análise Claude IA",
        "desc": "Enviar dados para Claude analisar propensão de compra",
        "saida": "score (0-100), análise textual",
    },
    {
        "num": 3,
        "nome": "Enviar Email (Resend)",
        "desc": "Email de boas-vindas personalizado",
        "saida": "email_sent: true/false",
    },
    {
        "num": 4,
        "nome": "Enviar SMS (Twilio)",
        "desc": "SMS de confirmação recebimento",
        "saida": "sms_sent: true/false",
    },
    {
        "num": 5,
        "nome": "Enviar WhatsApp (Official API)",
        "desc": "Mensagem template oficial via WhatsApp",
        "saida": "whatsapp_sent: true/false",
    },
    {
        "num": 6,
        "nome": "Atualizar CRM",
        "desc": "Salvar score, análise, status de mensagens",
        "saida": "lead updated in database",
    },
    {
        "num": 7,
        "nome": "Agendar Follow-up 24h",
        "desc": "Criar tarefa para follow-up automático em 24h",
        "saida": "followup scheduled",
    },
]

for step in ws_steps:
    p = doc.add_paragraph()
    p.add_run(f"{step['num']}. {step['nome']}: ").bold = True
    p.add_run(f"{step['desc']} → {step['saida']}")

# WF-02 até WF-05
followups_info = [
    {
        "wf": "WF-02",
        "titulo": "24h Follow-up (Sem Resposta)",
        "delay": "24 horas após lead criado",
        "condicao": "Sem interação registrada",
        "acao": "Email + SMS com pergunta de interesse",
    },
    {
        "wf": "WF-03",
        "titulo": "48h Follow-up (Sem Resposta)",
        "delay": "48 horas após lead criado",
        "condicao": "Sem resposta em 24h",
        "acao": "Email + SMS com oferta/incentivo",
    },
    {
        "wf": "WF-04",
        "titulo": "72h Follow-up (Sem Resposta)",
        "delay": "72 horas após lead criado",
        "condicao": "Sem resposta em 48h",
        "acao": "Email + WhatsApp com chamada para ação urgente",
    },
    {
        "wf": "WF-05",
        "titulo": "Lead Inativo 7 dias (Remarketing)",
        "delay": "7 dias sem contato",
        "condicao": "Sem interação há uma semana",
        "acao": "Email de reengajamento com novas opções",
    },
]

for info in followups_info:
    doc.add_heading(f"5.{['2','3','4','5'][followups_info.index(info)]} {info['wf']}: {info['titulo']}", level=2)

    p_delay = doc.add_paragraph()
    p_delay.add_run("Delay: ").bold = True
    p_delay.add_run(info['delay'])

    p_cond = doc.add_paragraph()
    p_cond.add_run("Condição: ").bold = True
    p_cond.add_run(info['condicao'])

    p_acao = doc.add_paragraph()
    p_acao.add_run("Ação: ").bold = True
    p_acao.add_run(info['acao'])

doc.add_page_break()

# WF-06
doc.add_heading("5.6 WF-06: Lead Qualificado (Score > 80)", level=2)

doc.add_paragraph("Acionado quando: Claude IA dá score > 80")
doc.add_paragraph()

qualified_steps = [
    "Verificar score no CRM (score >= 80)",
    "Enviar Email VIP com assunto destacado",
    "Enviar WhatsApp pessoal (não template)",
    "Atribuir prioridade ALTA no CRM",
    "Notificar melhor agente via email/SMS",
    "Agendar call com agente em 2 horas",
    "Registrar no CRM: lead_priority = HIGH",
]

for step in qualified_steps:
    doc.add_paragraph(step, style='List Bullet')

# WF-07
doc.add_heading("5.7 WF-07: Relatório Diário (9h da manhã)", level=2)

doc.add_paragraph("Acionado quando: Agendado para 9:00 AM todos os dias")
doc.add_paragraph()

relatorio_steps = [
    "Query: Contar leads criados ontem",
    "Query: Score médio dos leads",
    "Query: Leads com follow-up pendente",
    "Query: Taxa de conversão do mês",
    "Compilar em tabela HTML",
    "Enviar email para admin/gestor",
]

for step in relatorio_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_page_break()

# ============ 6. REMARKETING STRATEGY ============
doc.add_heading("6. REMARKETING DETALHADO", level=1)

doc.add_heading("6.1 Definição de Segmentos", level=2)

segmentos = {
    "Novo": "Lead criado há menos de 1 dia",
    "Quente": "Lead com interação nos últimos 3 dias",
    "Morno": "Lead sem interação há 3-7 dias",
    "Frio": "Lead sem interação há 7+ dias",
    "Qualificado": "Score IA >= 80",
    "Visitante": "Visitou imóvel mas não fechou",
    "Proposição": "Recebeu proposta de imóvel",
}

for seg, desc in segmentos.items():
    p = doc.add_paragraph()
    p.add_run(seg + ": ").bold = True
    p.add_run(desc)

doc.add_heading("6.2 Remarketing por Segmento", level=2)

remarketing = {
    "Leads Frios (7+ dias)": [
        "Email: 'Voltamos a falar sobre suas opções?'",
        "Oferta: 10% desconto em taxa de corretagem",
        "CTA: 'Agora quero ver as opções!'",
        "Frequency: 1x a cada 3 dias (máximo 3x)",
    ],
    "Leads Mortos (30+ dias)": [
        "Email: 'Ainda procurando imóvel?'",
        "Oferecimento: 'Novos imóveis que podem te interessar'",
        "CTA: 'Clique e veja as novidades'",
        "Frequency: 1x por semana",
    ],
    "Visitantes": [
        "Email: 'Gostou do imóvel que visitou?'",
        "Informação: 'Temos 5 imóveis similares disponíveis'",
        "CTA: 'Ver imóveis parecidos'",
        "Frequency: 1x após 24h da visita",
    ],
}

for segmento, acoes in remarketing.items():
    doc.add_heading(f"• {segmento}", level=3)
    for acao in acoes:
        doc.add_paragraph(acao, style='List Bullet')

doc.add_page_break()

# ============ 7. TIMELINE ============
doc.add_heading("7. TIMELINE DE IMPLEMENTAÇÃO", level=1)

doc.add_heading("Semana 1: Preparação e Configuração", level=2)

sem1 = [
    "Dia 1: Reunião kickoff, documentar especificações",
    "Dia 2: Setup servidor (n8n, PostgreSQL, domínio)",
    "Dia 3: Configurar credenciais (Claude, Resend, WhatsApp, Twilio)",
    "Dia 4: Customizar modelos Eloquent do CRM",
    "Dia 5: Criar dashboard e listagem de leads",
]

for item in sem1:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Semana 2: Integração e Automações", level=2)

sem2 = [
    "Dia 8: Implementar webhook CRM ← Jetimob",
    "Dia 9: Testar integração com form de teste",
    "Dia 10: Criar WF-01 (Welcome) no n8n",
    "Dia 11: Criar WF-02, WF-03, WF-04 (Follow-ups)",
    "Dia 12: Integrar Claude IA nos workflows",
]

for item in sem2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Semana 3: Remarketing e Testes", level=2)

sem3 = [
    "Dia 15: Criar WF-05 (Inativo 7 dias)",
    "Dia 16: Criar WF-06 (Lead Qualificado)",
    "Dia 17: Criar WF-07 (Relatório Diário)",
    "Dia 18: Testes end-to-end com dados reais",
    "Dia 19: Ajustes e otimizações",
]

for item in sem3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Semana 4: Deploy e Treinamento", level=2)

sem4 = [
    "Dia 22: Publicar em produção",
    "Dia 23: Monitoramento 24h",
    "Dia 24: Treinamento com time da YesMyHome",
    "Dia 25: Documentação final",
    "Dia 26: Suporte Go-live",
]

for item in sem4:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============ 8. DETALHES TÉCNICOS ============
doc.add_heading("8. DETALHES TÉCNICOS", level=1)

doc.add_heading("8.1 Claude API - Prompt de Análise", level=2)

prompt = """Você é um assistente de análise de leads imobiliários.

Analise o seguinte lead e retorne um JSON com:
- score (0-100): propensão de compra
- propensao (Alta/Média/Baixa)
- recomendacao (ação sugerida)
- tags (lista de categorias)

Lead:
- Nome: {name}
- Tipo de imóvel: {property_type}
- Tipo de operação: {operation_type}
- Orçamento: {budget}
- Localização: {location}
- Notas: {notes}

Retorne APENAS um JSON válido."""

prompt_para = doc.add_paragraph()
prompt_run = prompt_para.add_run(prompt)
prompt_run.font.name = 'Courier New'
prompt_run.font.size = Pt(9)

doc.add_heading("8.2 Credenciais Necessárias", level=2)

credenciais_table = doc.add_table(rows=7, cols=3)
credenciais_table.style = 'Table Grid'
credenciais_table.rows[0].cells[0].text = "Serviço"
credenciais_table.rows[0].cells[1].text = "Credencial"
credenciais_table.rows[0].cells[2].text = "Onde Configurar"

creds_data = [
    ("Claude API", "API Key", "console.anthropic.com"),
    ("Resend", "API Key", "resend.com/api-keys"),
    ("WhatsApp Business", "Token de Acesso", "business.facebook.com"),
    ("Twilio", "Account SID + Auth Token", "twilio.com/console"),
    ("PostgreSQL", "Connection String", ".env"),
]

for idx, (servico, cred, local) in enumerate(creds_data, start=1):
    credenciais_table.rows[idx].cells[0].text = servico
    credenciais_table.rows[idx].cells[1].text = cred
    credenciais_table.rows[idx].cells[2].text = local

doc.add_page_break()

# ============ 9. RISCOS E MITIGAÇÃO ============
doc.add_heading("9. RISCOS E MITIGAÇÃO", level=1)

riscos = [
    {
        "risco": "Email bloqueado por spam",
        "probabilidade": "Alta",
        "mitigacao": "Verificar domínio no Resend, configurar SPF/DKIM/DMARC",
    },
    {
        "risco": "WhatsApp Business API rejeitado",
        "probabilidade": "Média",
        "mitigacao": "Submeter para verificação com antecedência, ter documentação",
    },
    {
        "risco": "n8n cair/desconectar",
        "probabilidade": "Baixa",
        "mitigacao": "Health checks, alertas de uptime, backups automáticos",
    },
    {
        "risco": "Integração Jetimob falhar",
        "probabilidade": "Média",
        "mitigacao": "Testar webhook com curl, ter retry logic, logs detalhados",
    },
    {
        "risco": "Claude API indisponível",
        "probabilidade": "Muito Baixa",
        "mitigacao": "Fallback para score padrão (50), usar cache de análises",
    },
]

for risco_item in riscos:
    doc.add_heading(f"• {risco_item['risco']}", level=2)

    p_prob = doc.add_paragraph()
    p_prob.add_run("Probabilidade: ").bold = True
    p_prob.add_run(risco_item['probabilidade'])

    p_mit = doc.add_paragraph()
    p_mit.add_run("Mitigação: ").bold = True
    p_mit.add_run(risco_item['mitigacao'])

doc.add_page_break()

# ============ 10. PRÓXIMOS PASSOS ============
doc.add_heading("10. PRÓXIMOS PASSOS", level=1)

proximos = [
    "Aprovação deste assessment",
    "Confirmação de informações do Jetimob (access, formulários)",
    "Compra de número Twilio brasileiro",
    "Setup domínio para email (DNS records)",
    "Iniciar implementação Semana 1",
]

for item in proximos:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============ CONTATO ============
doc.add_heading("CONTATO", level=1)

contact_table = doc.add_table(rows=5, cols=2)
contact_table.style = 'Table Grid'
contact_table.rows[0].cells[0].text = "Empresa"
contact_table.rows[0].cells[1].text = "Tauri Tecnologia"
contact_table.rows[1].cells[0].text = "Dev Lead"
contact_table.rows[1].cells[1].text = "Romulo"
contact_table.rows[2].cells[0].text = "Email"
contact_table.rows[2].cells[1].text = "romulo@tauritecnologia.com.br"
contact_table.rows[3].cells[0].text = "WhatsApp"
contact_table.rows[3].cells[1].text = "(41) 98519-7035"
contact_table.rows[4].cells[0].text = "Data"
contact_table.rows[4].cells[1].text = "30 de julho de 2026"

# Salvar
doc.save("ASSESSMENT_TECNICO_YESMYHOME.docx")
print("✅ Assessment Técnico criado com sucesso!")
print("📄 Arquivo: ASSESSMENT_TECNICO_YESMYHOME.docx")
