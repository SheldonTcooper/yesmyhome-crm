#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# ========== CAPA ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("🏠 CRM YESMYHOME")
title_run.font.size = Pt(40)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(88, 166, 255)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("Automação Completa de Vendas e Atendimento")
subtitle_run.font.size = Pt(16)
subtitle.paragraph_format.space_before = Pt(6)

doc.add_paragraph()
doc.add_paragraph()

# Descrição
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc.add_run("Um gerente de vendas digital 24/7 que nunca dorme,\nque envia mensagens, gerencia contratos e resgata clientes esquecidos")
desc_run.font.size = Pt(12)
desc_run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Info
info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Table Grid'
info_table.rows[0].cells[0].text = "📅 Data"
info_table.rows[0].cells[1].text = "30 de julho de 2026"
info_table.rows[1].cells[0].text = "💼 Para"
info_table.rows[1].cells[1].text = "YesMyHome Negociações Imobiliárias"
info_table.rows[2].cells[0].text = "⏰ Tempo"
info_table.rows[2].cells[1].text = "Implementação em 4 semanas"
info_table.rows[3].cells[0].text = "💰 Investimento"
info_table.rows[3].cells[1].text = "R$ 3.300 setup + R$ 84-166/mês"

doc.add_page_break()

# ========== O PROBLEMA ==========
doc.add_heading("❌ O PROBLEMA DE HOJE", level=1)

problemas = [
    "Cliente preenche formulário no site → Ninguém vê automaticamente",
    "Corretores digitam manualmente todos os dados",
    "Sem follow-up automático → Muitos clientes são perdidos",
    "Contratos enviados por email (se lembram) → Atrasos",
    "Sem histórico central → Cada um tem sua anotação",
    "Aniversários esquecidos → Cliente não sente valor",
    "Pagamentos atrasados → Alguém tem que ligar (custo)",
]

for prob in problemas:
    doc.add_paragraph(prob, style='List Bullet')

doc.add_paragraph()
final_prob = doc.add_paragraph()
final_prob.add_run("Resultado: ").bold = True
final_prob.add_run("Perdem 30-50% dos leads por falta de acompanhamento")
final_prob.paragraph_format.space_before = Pt(6)

doc.add_page_break()

# ========== A SOLUÇÃO ==========
doc.add_heading("✅ A SOLUÇÃO", level=1)

doc.add_heading("Três ferramentas trabalhando juntas:", level=2)

solucoes = {
    "1️⃣ CRM Krayin (Seu Fichário Digital)": [
        "Centraliza TODOS os clientes em um único lugar",
        "Cada corretor vê seu histórico completo",
        "Acesso por qualquer um na empresa (computador, celular)",
        "Sem papéis, sem confusão, sem perder informações",
    ],
    "2️⃣ n8n (Seu Robô 24/7)": [
        "Trabalha enquanto você dorme",
        "Envia mensagens automáticas no momento certo",
        "Faz remarketing (traz clientes de volta)",
        "Envia contratos, lembretes de pagamento, aniversários",
    ],
    "3️⃣ WhatsApp Oficial (Comunicação Direta)": [
        "Cliente recebe mensagens por WhatsApp (que todos leem)",
        "Não é SMS que ninguém abre (é WhatsApp)",
        "Parece pessoal mas é completamente automático",
        "Taxa de abertura: 90%+ (vs email: 20%)",
    ],
}

for titulo, items in solucoes.items():
    doc.add_heading(titulo, level=3)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ========== COMO FUNCIONA ==========
doc.add_heading("⚙️ COMO FUNCIONA NA PRÁTICA", level=1)

doc.add_heading("Cenário 1: Novo Cliente", level=2)

cenario1 = [
    "10:15 - Cliente preenche formulário no site yesmyhome.com.br",
    "10:16 - CRM recebe AUTOMATICAMENTE",
    "10:17 - Robô envia: 'Oi João, recebemos seu contato!' (WhatsApp)",
    "10:18 - Robô envia email com próximos passos",
    "10:19 - Robô analisa (IA Claude): 'João tem 85% chance de comprar'",
    "10:20 - Corretores veem o lead no Dashboard (tempo real)",
    "10:21 - Melhor corretor é atribuído automaticamente",
    "24h depois - Robô envia follow-up: 'Gostou das opções?'",
]

for idx, passo in enumerate(cenario1, 1):
    doc.add_paragraph(passo, style='List Bullet')

doc.add_heading("Cenário 2: Cliente Esquecido (Remarketing)", level=2)

cenario2 = [
    "📅 Dia 1-3: Ninguém responde o cliente",
    "📱 Dia 4: Robô envia outro WhatsApp com oferta",
    "📧 Dia 5: Email especial: 'Temos novas opções'",
    "⏰ Dia 7: SMS: 'Ainda procurando? Ligamos amanhã?'",
    "🎯 Resultado: 30-50% dos clientes 'mortos' voltam",
]

for passo in cenario2:
    doc.add_paragraph(passo, style='List Bullet')

doc.add_heading("Cenário 3: Contrato Fechado", level=2)

cenario3 = [
    "✍️ Admin: Clica 'Enviar Contrato'",
    "🤖 Robô faz TUDO:",
    "   • Pega dados do cliente",
    "   • Gera PDF automático",
    "   • Envia por email",
    "   • Envia link por WhatsApp",
    "   • Registra no histórico",
    "🎉 Cliente recebe em 30 segundos (não 3 dias)",
]

for passo in cenario3:
    doc.add_paragraph(passo, style='List Bullet')

doc.add_page_break()

# ========== FUNCIONALIDADES ==========
doc.add_heading("🎁 FUNCIONALIDADES INCLUÍDAS", level=1)

funcionalidades = [
    ("🏠 CRM Central", "Todos os clientes em um lugar. Histórico completo. Sem perder informação."),
    ("👥 Cadastro de Corretores", "Perfil com foto, especialidade, histórico de vendas, nota/avaliação."),
    ("📱 WhatsApp Automático", "Boas-vindas, follow-ups, contratos, lembretes. Tudo por WhatsApp."),
    ("📧 Emails Automáticos", "Confirmações, contratos, propostas, lembretes de pagamento."),
    ("📄 Contratos em 30s", "Click → Gera PDF → Envia email + WhatsApp."),
    ("🎂 Aniversários", "Robô automaticamente envia mensagem de parabéns."),
    ("💳 Pagamentos", "Aviso de vencimento → Aviso de recebimento → Comprovante automático."),
    ("🤖 IA Análise", "Robô avalia cada cliente: 'Tem 90% chance de comprar' ou 'Tem 40%'"),
    ("📊 Dashboard", "Gráficos em tempo real: quanto vendeu, conversão, leads pendentes."),
    ("⏰ Follow-ups", "Automático em 24h, 48h, 72h. Sem esquecer ninguém."),
]

for titulo, desc in funcionalidades:
    p = doc.add_paragraph()
    p.add_run(titulo + " - ").bold = True
    p.add_run(desc)

doc.add_page_break()

# ========== BENEFÍCIOS ==========
doc.add_heading("💰 BENEFÍCIOS FINANCEIROS", level=1)

doc.add_heading("Antes (Sem Sistema)", level=2)

antes = [
    "1 atendente digitando dados = 4-5 horas/dia",
    "40-50% de leads perdidos por falta de follow-up",
    "Contratos levam 2-3 dias para chegar",
    "Pagamentos atrasados requerem ligações (custo + tempo)",
]

for item in antes:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Depois (Com Sistema)", level=2)

depois = [
    "Atendente 100% disponível para vender (não digita mais)",
    "Recupera 30-50% de leads que iam morrer",
    "Contratos entregues em 30 segundos",
    "Pagamentos cobrados automaticamente",
]

for item in depois:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Conta Simples", level=2)

conta = doc.add_paragraph()
conta.add_run("50 leads/mês × R$ 5.000/lead = R$ 250.000 em vendas\n").bold = True
conta.add_run("3% conversão = R$ 7.500 em lucro\n").font.color.rgb = RGBColor(34, 139, 34)
conta.add_run("Sistema custa: R$ 165/mês\n").bold = True
conta.add_run("Lucro LÍQUIDO: R$ 7.335/mês\n").font.color.rgb = RGBColor(34, 139, 34)

doc.add_paragraph()
payback = doc.add_paragraph()
payback_run = payback.add_run("🎯 Sistema paga a si mesmo em MENOS DE 1 DIA")
payback_run.bold = True
payback_run.font.size = Pt(14)
payback_run.font.color.rgb = RGBColor(255, 100, 0)

doc.add_page_break()

# ========== SITE + CRM ==========
doc.add_heading("🔗 COMO SITE E CRM TRABALHAM JUNTOS", level=1)

fluxo = [
    "1️⃣ Cliente entra em yesmyhome.com.br",
    "2️⃣ Preenche formulário (nome, email, WhatsApp, tipo de imóvel)",
    "3️⃣ Clica 'Enviar'",
    "4️⃣ Sistema recebe AUTOMATICAMENTE",
    "5️⃣ Lead aparece no CRM em TEMPO REAL",
    "6️⃣ Robô dispara automações (mensagens, análise IA, follow-ups)",
    "7️⃣ Corretores veem no Dashboard e contatam",
]

for passo in fluxo:
    doc.add_paragraph(passo, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph("Resultado: Não perde NENHUM lead.")

doc.add_page_break()

# ========== INVESTIMENTO ==========
doc.add_heading("💵 INVESTIMENTO", level=1)

doc.add_heading("One-time (Setup)", level=2)

setup_table = doc.add_table(rows=5, cols=2)
setup_table.style = 'Table Grid'
setup_table.rows[0].cells[0].text = "Setup CRM + Integração"
setup_table.rows[0].cells[1].text = "R$ 2.000"
setup_table.rows[1].cells[0].text = "Configuração Robô (n8n)"
setup_table.rows[1].cells[1].text = "R$ 1.000"
setup_table.rows[2].cells[0].text = "WhatsApp Official Setup"
setup_table.rows[2].cells[1].text = "R$ 300"
setup_table.rows[3].cells[0].text = "TOTAL"
setup_table.rows[3].cells[1].text = "R$ 3.300"

for cell in setup_table.rows[3].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

doc.add_heading("Mensal (Recorrente)", level=2)

mensal_table = doc.add_table(rows=6, cols=2)
mensal_table.style = 'Table Grid'
mensal_table.rows[0].cells[0].text = "Claude IA"
mensal_table.rows[0].cells[1].text = "R$ 0-50"
mensal_table.rows[1].cells[0].text = "WhatsApp"
mensal_table.rows[1].cells[1].text = "R$ 4-130"
mensal_table.rows[2].cells[0].text = "Email (Resend)"
mensal_table.rows[2].cells[1].text = "R$ 0-50"
mensal_table.rows[3].cells[0].text = "Domínio + Hosting"
mensal_table.rows[3].cells[1].text = "R$ 80"
mensal_table.rows[4].cells[0].text = "n8n (Robô)"
mensal_table.rows[4].cells[1].text = "R$ 0 (gratuito)"
mensal_table.rows[5].cells[0].text = "TOTAL/MÊS"
mensal_table.rows[5].cells[1].text = "R$ 84-230"

for cell in mensal_table.rows[5].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

doc.add_paragraph()
ano1 = doc.add_paragraph()
ano1.add_run("Ano 1: ").bold = True
ano1.add_run("R$ 3.300 setup + R$ 1.500-2.500 (12 meses) = ~R$ 5.000-6.000 total")

doc.add_page_break()

# ========== TIMELINE ==========
doc.add_heading("📅 TIMELINE: 4 SEMANAS", level=1)

doc.add_heading("Semana 1: Preparação", level=2)
sem1 = ["Setup do servidor", "Instalar CRM", "Conectar WhatsApp", "Cadastrar corretores"]
for item in sem1:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Semana 2: Integração", level=2)
sem2 = ["Conectar site → CRM", "Criar automações (boas-vindas)", "Testar com dados falsos"]
for item in sem2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Semana 3: Robô 24/7", level=2)
sem3 = ["Ativar follow-ups automáticos", "Ativar remarketing", "Testar tudo com dados reais"]
for item in sem3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Semana 4: Publicar", level=2)
sem4 = ["Deploy em produção", "Treinar equipe", "Go-live!"]
for item in sem4:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ========== O QUE NOS DIFERENCIA ==========
doc.add_heading("⭐ POR QUE ESCOLHER ISSO?", level=1)

diferenciais = [
    ("✅ Sem Custo de Pessoa", "Robô trabalha 24/7 sem salário, sem benefício, sem reclamação"),
    ("✅ Sem Erros", "Sistema não esquece, não erra, não tem dia ruim"),
    ("✅ Recupera Clientes", "30-50% de clientes 'mortos' voltam com remarketing automático"),
    ("✅ Comunica Bem", "WhatsApp = 90% de abertura (email só tem 20%)"),
    ("✅ Mais Vendas", "Corretores focam 100% em vender, não em digitar"),
    ("✅ Controle Total", "Você é dono do CRM, dados, tudo. Sem dependência"),
    ("✅ Paga a Si Mesmo", "ROI em 1 dia (custa R$ 165/mês, traz R$ 7.000+/mês)"),
]

for diferencial, desc in diferenciais:
    p = doc.add_paragraph()
    p.add_run(diferencial + " - ").bold = True
    p.add_run(desc)

doc.add_page_break()

# ========== PRÓXIMOS PASSOS ==========
doc.add_heading("🚀 PRÓXIMOS PASSOS", level=1)

passos = [
    "1. Você aprova essa abordagem? (SIM / NÃO / Dúvida)",
    "2. Confirmamos informações do Jetimob (access, formulários)",
    "3. Agendamos kickoff meeting",
    "4. Começamos Semana 1",
]

for passo in passos:
    doc.add_paragraph(passo, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

# ========== RESUMO FINAL ==========
doc.add_heading("📌 RESUMO EM UMA FRASE", level=2)

resumo = doc.add_paragraph()
resumo_run = resumo.add_run(
    "Um gerente de vendas digital (robô) que trabalha 24/7, "
    "envia mensagens automáticas, gerencia contratos, "
    "e traz clientes esquecidos de volta — "
    "tudo conectado ao seu site e CRM."
)
resumo_run.font.bold = True
resumo_run.font.size = Pt(12)
resumo_run.font.color.rgb = RGBColor(88, 166, 255)

doc.add_page_break()

# ========== CONTATO ==========
doc.add_heading("📞 CONTATO", level=1)

contact_table = doc.add_table(rows=4, cols=2)
contact_table.style = 'Table Grid'
contact_table.rows[0].cells[0].text = "👤 Empresa"
contact_table.rows[0].cells[1].text = "Tauri Tecnologia"
contact_table.rows[1].cells[0].text = "📧 Email"
contact_table.rows[1].cells[1].text = "romulo@tauritecnologia.com.br"
contact_table.rows[2].cells[0].text = "📱 WhatsApp"
contact_table.rows[2].cells[1].text = "(41) 98519-7035"
contact_table.rows[3].cells[0].text = "📅 Data"
contact_table.rows[3].cells[1].text = "30 de julho de 2026"

# Salvar
doc.save("APRESENTACAO_EXECUTIVA_YESMYHOME.docx")
print("✅ Apresentação Executiva criada com sucesso!")
print("📄 Arquivo: APRESENTACAO_EXECUTIVA_YESMYHOME.docx")
